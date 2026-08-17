"""
OLT MOP Automation - Streamlit Web Application
Compare FIO (Excel), EWP (Image), and MOP (Docx) files between two sites
"""

import streamlit as st
import pandas as pd
import numpy as np
import os
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import logging
from dataclasses import dataclass, asdict
from enum import Enum
import tempfile
import base64
from io import BytesIO

# File processing libraries
import openpyxl
from docx import Document
from docx.table import Table
from PIL import Image
import cv2
import numpy as np

# Try to import pytesseract, but handle gracefully
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None

# Report generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table as RLTable, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Set page config
st.set_page_config(
    page_title="OLT MOP Automation Tool",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #0066cc;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .diff-critical {
        background-color: #ffcccc;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #ff0000;
        margin-bottom: 0.5rem;
    }
    .diff-warning {
        background-color: #ffe6cc;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #ff9900;
        margin-bottom: 0.5rem;
    }
    .diff-info {
        background-color: #d4edda;
        padding: 0.5rem;
        border-radius: 5px;
        border-left: 5px solid #28a745;
        margin-bottom: 0.5rem;
    }
    .stButton > button {
        width: 100%;
        font-weight: bold;
    }
    .upload-section {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #dee2e6;
        margin-bottom: 1rem;
    }
    .file-uploaded {
        background-color: #d4edda;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        color: #155724;
        font-size: 0.85rem;
        display: inline-block;
        margin-top: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileType(Enum):
    FIO = "fio"
    EWP = "ewp"
    MOP = "mop"


@dataclass
class SiteData:
    """Data structure for a single site's configuration"""
    site_name: str
    fio_data: Dict[str, Any]
    ewp_data: Dict[str, Any]
    mop_data: Dict[str, Any]
    fio_file: str
    ewp_file: str
    mop_file: str
    
    def to_dict(self) -> Dict:
        return {
            'site_name': self.site_name,
            'fio_data': self.fio_data,
            'ewp_data': self.ewp_data,
            'mop_data': self.mop_data,
            'fio_file': self.fio_file,
            'ewp_file': self.ewp_file,
            'mop_file': self.mop_file
        }


@dataclass
class Difference:
    """Represents a difference between two files"""
    category: str
    field: str
    site1_value: Any
    site2_value: Any
    source1: str
    source2: str
    severity: str
    description: str
    
    def to_dict(self) -> Dict:
        return asdict(self)


class FIOProcessor:
    """Processes FIO Excel files"""
    
    def __init__(self):
        self.required_columns = ['OLT_ID', 'SLOT', 'PORT', 'SERVICE', 'VLAN', 'BANDWIDTH']
        self.data = {}
    
    def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract structured data from FIO Excel file"""
        try:
            logger.info(f"Processing FIO file: {filename}")
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            df = pd.read_excel(tmp_path, sheet_name=0)
            os.unlink(tmp_path)
            
            # Clean column names
            df.columns = df.columns.str.strip().str.upper()
            
            # Extract OLT configuration
            config = {
                'olt_connections': [],
                'services': [],
                'vlans': set(),
                'slots': set(),
                'ports': set(),
                'summary': {}
            }
            
            # Map columns to expected format
            column_mapping = self._map_columns(df.columns)
            
            for _, row in df.iterrows():
                connection = {}
                for col, mapped_col in column_mapping.items():
                    if col in df.columns:
                        val = row[col]
                        if pd.notna(val):
                            if isinstance(val, (int, float)):
                                val = str(int(val)) if val.is_integer() else str(val)
                            connection[mapped_col] = str(val)
                
                if connection:
                    config['olt_connections'].append(connection)
                    if 'VLAN' in connection:
                        config['vlans'].add(connection['VLAN'])
                    if 'SLOT' in connection:
                        config['slots'].add(connection['SLOT'])
                    if 'PORT' in connection:
                        config['ports'].add(connection['PORT'])
                    if 'SERVICE' in connection:
                        config['services'].append(connection['SERVICE'])
            
            # Generate summary
            config['summary'] = {
                'total_connections': len(config['olt_connections']),
                'unique_vlans': len(config['vlans']),
                'unique_slots': len(config['slots']),
                'unique_ports': len(config['ports']),
                'unique_services': len(set(config['services']))
            }
            
            self.data = config
            return config
            
        except Exception as e:
            logger.error(f"Error processing FIO file: {e}")
            raise
    
    def _map_columns(self, columns: List[str]) -> Dict[str, str]:
        """Map actual column names to expected format"""
        mapping = {}
        for col in columns:
            col_upper = col.strip().upper()
            for expected in self.required_columns:
                if expected in col_upper or col_upper in expected:
                    mapping[col] = expected
                    break
        return mapping


class EWPProcessor:
    """Processes EWP Image files using OCR with fallback"""
    
    def __init__(self):
        self.data = {}
        self.tesseract_config = '--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ-/:.'
    
    def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract structured data from EWP image using OCR or fallback"""
        try:
            logger.info(f"Processing EWP file: {filename}")
            
            # Check if Tesseract is available
            if not TESSERACT_AVAILABLE:
                return self._fallback_processing(file_content, filename)
            
            # Save to temporary file
            suffix = Path(filename).suffix.lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                # Try OpenCV
                img = cv2.imread(tmp_path)
                if img is None:
                    # Fallback to PIL
                    img = Image.open(tmp_path)
                    img = np.array(img)
                
                # Preprocess image
                processed_img = self._preprocess_image(img)
                
                # Extract text using OCR
                text = pytesseract.image_to_string(processed_img, config=self.tesseract_config)
                
                # Parse the extracted text
                parsed_data = self._parse_ewp_text(text)
                
            except Exception as e:
                logger.warning(f"OCR processing failed, using fallback: {e}")
                parsed_data = self._fallback_processing(file_content, filename)
            finally:
                try:
                    os.unlink(tmp_path)
                except:
                    pass
            
            self.data = parsed_data
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error processing EWP file: {e}")
            return self._fallback_processing(file_content, filename)
    
    def _fallback_processing(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback processing when OCR is not available"""
        logger.info("Using fallback EWP processing")
        
        parsed = {
            'olt_configurations': [],
            'ports': [],
            'vlans': [],
            'services': [],
            'raw_text': f"EWP Image: {filename}",
            'filename': filename,
            'file_size': len(file_content),
            'ocr_available': False
        }
        
        # Try to extract info from filename
        name_parts = filename.replace('_', ' ').replace('-', ' ').split()
        for part in name_parts:
            part_upper = part.upper()
            if 'OLT' in part_upper:
                olt_match = re.search(r'OLT[-\s]*(\d+)', part_upper)
                if olt_match:
                    parsed['olt_configurations'].append({'OLT': olt_match.group(1)})
                else:
                    parsed['olt_configurations'].append({'identifier': part})
            if 'VLAN' in part_upper:
                vlan_match = re.search(r'VLAN[-\s]*(\d+)', part_upper)
                if vlan_match:
                    parsed['vlans'].append(vlan_match.group(1))
            if 'PORT' in part_upper:
                port_match = re.search(r'PORT[-\s]*(\d+)', part_upper)
                if port_match:
                    parsed['ports'].append(port_match.group(1))
        
        parsed['summary'] = {
            'total_olt_configs': len(parsed['olt_configurations']),
            'total_ports': len(parsed['ports']),
            'total_vlans': len(set(parsed['vlans'])),
            'total_services': len(set(parsed['services'])),
            'ocr_available': False
        }
        
        return parsed
    
    def _preprocess_image(self, img: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR accuracy"""
        if len(img.shape) == 3:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        else:
            gray = img
        
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        denoised = cv2.fastNlMeansDenoising(thresh)
        
        return denoised
    
    def _parse_ewp_text(self, text: str) -> Dict[str, Any]:
        """Parse OCR extracted text into structured data"""
        lines = text.split('\n')
        parsed = {
            'olt_configurations': [],
            'ports': [],
            'vlans': [],
            'services': [],
            'raw_text': text,
            'ocr_available': True
        }
        
        current_config = {}
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if 'OLT' in line.upper() or 'PON' in line.upper():
                if current_config:
                    parsed['olt_configurations'].append(current_config)
                current_config = {'section': line}
            elif ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().upper()
                value = value.strip()
                if key:
                    current_config[key] = value
                    if 'VLAN' in key:
                        parsed['vlans'].append(value)
                    elif 'PORT' in key:
                        parsed['ports'].append(value)
                    elif 'SERVICE' in key or 'SERV' in key:
                        parsed['services'].append(value)
            else:
                if self._is_olt_identifier(line):
                    if current_config:
                        parsed['olt_configurations'].append(current_config)
                    current_config = {'identifier': line}
        
        if current_config:
            parsed['olt_configurations'].append(current_config)
        
        parsed['summary'] = {
            'total_olt_configs': len(parsed['olt_configurations']),
            'total_ports': len(parsed['ports']),
            'total_vlans': len(set(parsed['vlans'])),
            'total_services': len(set(parsed['services'])),
            'ocr_available': True
        }
        
        return parsed
    
    def _is_olt_identifier(self, text: str) -> bool:
        """Check if text looks like an OLT identifier"""
        patterns = [
            r'OLT-\d+',
            r'PON-\d+',
            r'^\d+/\d+/\d+',
            r'^\d+-\d+-\d+',
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False


class MOPProcessor:
    """Processes MOP Docx files"""
    
    def __init__(self):
        self.data = {}
    
    def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract structured data from MOP Docx file"""
        try:
            logger.info(f"Processing MOP file: {filename}")
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            doc = Document(tmp_path)
            os.unlink(tmp_path)
            
            parsed = {
                'sections': [],
                'olt_configurations': [],
                'parameters': {},
                'tables': [],
                'summary': {}
            }
            
            # Process paragraphs
            current_section = {'title': '', 'content': []}
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                if para.style.name and 'Heading' in para.style.name:
                    if current_section['title'] or current_section['content']:
                        parsed['sections'].append(current_section)
                    current_section = {'title': text, 'content': []}
                else:
                    current_section['content'].append(text)
                    
                    if 'OLT' in text.upper() or 'PON' in text.upper():
                        config = self._parse_olt_config(text)
                        if config:
                            parsed['olt_configurations'].append(config)
                    
                    if ':' in text:
                        key, value = text.split(':', 1)
                        parsed['parameters'][key.strip()] = value.strip()
            
            if current_section['title'] or current_section['content']:
                parsed['sections'].append(current_section)
            
            # Process tables
            for table in doc.tables:
                table_data = self._parse_table(table)
                if table_data:
                    parsed['tables'].append(table_data)
                    self._extract_olt_from_table(table_data, parsed)
            
            parsed['summary'] = {
                'total_sections': len(parsed['sections']),
                'total_tables': len(parsed['tables']),
                'total_olt_configs': len(parsed['olt_configurations']),
                'total_parameters': len(parsed['parameters'])
            }
            
            self.data = parsed
            return parsed
            
        except Exception as e:
            logger.error(f"Error processing MOP file: {e}")
            raise
    
    def _parse_table(self, table: Table) -> List[List[str]]:
        """Parse a docx table into a list of rows"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                row_data.append(text)
            if any(row_data):
                table_data.append(row_data)
        return table_data
    
    def _extract_olt_from_table(self, table_data: List[List[str]], parsed: Dict):
        """Extract OLT configuration from table data"""
        if not table_data:
            return
        
        headers = table_data[0] if table_data else []
        olt_columns = ['OLT', 'SLOT', 'PORT', 'VLAN', 'SERVICE', 'BANDWIDTH', 'SPEED']
        
        col_indices = {}
        for idx, header in enumerate(headers):
            header_upper = header.upper()
            for olt_col in olt_columns:
                if olt_col in header_upper or header_upper in olt_col:
                    col_indices[olt_col] = idx
                    break
        
        for row in table_data[1:]:
            if len(row) < len(headers):
                continue
            
            config = {}
            for col, idx in col_indices.items():
                if idx < len(row) and row[idx]:
                    config[col] = row[idx].strip()
            
            if config:
                parsed['olt_configurations'].append(config)
    
    def _parse_olt_config(self, text: str) -> Optional[Dict]:
        """Parse OLT configuration from text"""
        config = {}
        parts = text.split()
        for part in parts:
            if '=' in part:
                key, value = part.split('=', 1)
                config[key.strip()] = value.strip()
            elif ':' in part:
                key, value = part.split(':', 1)
                config[key.strip()] = value.strip()
        
        return config if config else None


class DifferenceAnalyzer:
    """Analyzes differences between site data"""
    
    def __init__(self):
        self.differences = []
    
    def analyze(self, site1: SiteData, site2: SiteData) -> List[Difference]:
        """Analyze differences between two sites"""
        self.differences = []
        
        self._compare_fio_data(site1, site2)
        self._compare_ewp_data(site1, site2)
        self._compare_mop_data(site1, site2)
        self._cross_reference_data(site1, site2)
        
        return self.differences
    
    def _compare_fio_data(self, site1: SiteData, site2: SiteData):
        """Compare FIO data between sites"""
        fio1 = site1.fio_data.get('olt_connections', [])
        fio2 = site2.fio_data.get('olt_connections', [])
        
        if len(fio1) != len(fio2):
            self.differences.append(Difference(
                category='FIO',
                field='total_connections',
                site1_value=len(fio1),
                site2_value=len(fio2),
                source1=site1.fio_file,
                source2=site2.fio_file,
                severity='warning',
                description=f'Different number of OLT connections'
            ))
        
        fields_to_compare = ['OLT_ID', 'SLOT', 'PORT', 'SERVICE', 'VLAN', 'BANDWIDTH']
        for i, (conn1, conn2) in enumerate(zip(fio1, fio2)):
            for field in fields_to_compare:
                val1 = conn1.get(field, '')
                val2 = conn2.get(field, '')
                if val1 != val2:
                    self.differences.append(Difference(
                        category='FIO',
                        field=field,
                        site1_value=val1,
                        site2_value=val2,
                        source1=site1.fio_file,
                        source2=site2.fio_file,
                        severity='critical' if field in ['OLT_ID', 'SLOT', 'PORT'] else 'warning',
                        description=f'Connection {i+1}: {field} mismatch'
                    ))
    
    def _compare_ewp_data(self, site1: SiteData, site2: SiteData):
        """Compare EWP data between sites"""
        ewp1 = site1.ewp_data.get('olt_configurations', [])
        ewp2 = site2.ewp_data.get('olt_configurations', [])
        
        for i, (cfg1, cfg2) in enumerate(zip(ewp1, ewp2)):
            for key in set(cfg1.keys()) | set(cfg2.keys()):
                val1 = cfg1.get(key, '')
                val2 = cfg2.get(key, '')
                if val1 != val2:
                    self.differences.append(Difference(
                        category='EWP',
                        field=key,
                        site1_value=val1,
                        site2_value=val2,
                        source1=site1.ewp_file,
                        source2=site2.ewp_file,
                        severity='critical' if key in ['OLT', 'PON'] else 'warning',
                        description=f'Configuration {i+1}: {key} mismatch'
                    ))
    
    def _compare_mop_data(self, site1: SiteData, site2: SiteData):
        """Compare MOP data between sites"""
        mop1 = site1.mop_data.get('olt_configurations', [])
        mop2 = site2.mop_data.get('olt_configurations', [])
        
        for i, (cfg1, cfg2) in enumerate(zip(mop1, mop2)):
            for key in set(cfg1.keys()) | set(cfg2.keys()):
                val1 = cfg1.get(key, '')
                val2 = cfg2.get(key, '')
                if val1 != val2:
                    self.differences.append(Difference(
                        category='MOP',
                        field=key,
                        site1_value=val1,
                        site2_value=val2,
                        source1=site1.mop_file,
                        source2=site2.mop_file,
                        severity='critical' if key in ['OLT', 'SLOT', 'PORT'] else 'warning',
                        description=f'Configuration {i+1}: {key} mismatch'
                    ))
        
        sections1 = site1.mop_data.get('sections', [])
        sections2 = site2.mop_data.get('sections', [])
        if len(sections1) != len(sections2):
            self.differences.append(Difference(
                category='MOP',
                field='sections_count',
                site1_value=len(sections1),
                site2_value=len(sections2),
                source1=site1.mop_file,
                source2=site2.mop_file,
                severity='info',
                description='Different number of sections in MOP documents'
            ))
    
    def _cross_reference_data(self, site1: SiteData, site2: SiteData):
        """Cross-reference data between different file types"""
        fio_olts = set()
        for conn in site1.fio_data.get('olt_connections', []):
            if 'OLT_ID' in conn:
                fio_olts.add(conn['OLT_ID'])
        
        ewp_olts = set()
        for cfg in site1.ewp_data.get('olt_configurations', []):
            for value in cfg.values():
                if isinstance(value, str) and ('OLT' in value or 'PON' in value):
                    ewp_olts.add(value)
        
        missing_in_ewp = fio_olts - ewp_olts
        if missing_in_ewp:
            self.differences.append(Difference(
                category='Cross-reference',
                field='FIO_to_EWP',
                site1_value=list(fio_olts),
                site2_value=list(ewp_olts),
                source1=site1.fio_file,
                source2=site1.ewp_file,
                severity='critical',
                description=f'FIO OLT IDs missing in EWP: {missing_in_ewp}'
            ))


class ReportGenerator:
    """Generates reports from difference analysis"""
    
    def __init__(self):
        self.differences = []
        self.site1_name = ''
        self.site2_name = ''
    
    def generate_summary(self, differences: List[Difference], site1_name: str, site2_name: str) -> Dict:
        """Generate a summary of differences"""
        self.differences = differences
        self.site1_name = site1_name
        self.site2_name = site2_name
        
        summary = {
            'total_differences': len(differences),
            'critical_count': 0,
            'warning_count': 0,
            'info_count': 0,
            'categories': {},
            'fields': {},
            'site1_name': site1_name,
            'site2_name': site2_name,
            'generated_at': datetime.now().isoformat()
        }
        
        for diff in differences:
            if diff.severity == 'critical':
                summary['critical_count'] += 1
            elif diff.severity == 'warning':
                summary['warning_count'] += 1
            else:
                summary['info_count'] += 1
            
            if diff.category not in summary['categories']:
                summary['categories'][diff.category] = 0
            summary['categories'][diff.category] += 1
            
            if diff.field not in summary['fields']:
                summary['fields'][diff.field] = 0
            summary['fields'][diff.field] += 1
        
        return summary
    
    def generate_excel_report(self, differences: List[Difference], site1_name: str, site2_name: str) -> BytesIO:
        """Generate an Excel report as bytes"""
        data = []
        for diff in differences:
            data.append({
                'Category': diff.category,
                'Field': diff.field,
                f'{site1_name}_Value': diff.site1_value,
                f'{site2_name}_Value': diff.site2_value,
                'Source_1': diff.source1,
                'Source_2': diff.source2,
                'Severity': diff.severity,
                'Description': diff.description
            })
        
        df = pd.DataFrame(data)
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Differences')
            
            summary_df = pd.DataFrame([
                ['Total Differences', len(differences)],
                ['Critical', sum(1 for d in differences if d.severity == 'critical')],
                ['Warning', sum(1 for d in differences if d.severity == 'warning')],
                ['Info', sum(1 for d in differences if d.severity == 'info')]
            ], columns=['Metric', 'Value'])
            summary_df.to_excel(writer, index=False, sheet_name='Summary')
        
        output.seek(0)
        return output
    
    def generate_pdf_report(self, differences: List[Difference], summary: Dict) -> BytesIO:
        """Generate a PDF report as bytes"""
        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=landscape(letter))
        styles = getSampleStyleSheet()
        story = []
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph("OLT MOP Automation - Difference Report", title_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph(f"Generated: {summary['generated_at']}", styles['Normal']))
        story.append(Paragraph(f"Site 1: {summary['site1_name']}", styles['Normal']))
        story.append(Paragraph(f"Site 2: {summary['site2_name']}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Summary Statistics", styles['Heading2']))
        stats_data = [
            ['Total Differences', str(summary['total_differences'])],
            ['Critical', str(summary['critical_count'])],
            ['Warning', str(summary['warning_count'])],
            ['Info', str(summary['info_count'])]
        ]
        stats_table = RLTable(stats_data, colWidths=[2*inch, 1*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("Detailed Differences", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        if not differences:
            story.append(Paragraph("No differences found!", styles['Normal']))
        else:
            table_data = [['#', 'Category', 'Field', f'{summary["site1_name"]}', f'{summary["site2_name"]}', 'Severity', 'Description']]
            for i, diff in enumerate(differences, 1):
                table_data.append([
                    str(i),
                    diff.category,
                    diff.field,
                    str(diff.site1_value)[:50],
                    str(diff.site2_value)[:50],
                    diff.severity.upper(),
                    diff.description[:100]
                ])
            
            col_widths = [0.5*inch, 1*inch, 1.2*inch, 1.5*inch, 1.5*inch, 1*inch, 2*inch]
            detail_table = RLTable(table_data, colWidths=col_widths, repeatRows=1)
            
            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ])
            
            for i, diff in enumerate(differences, 1):
                if diff.severity == 'critical':
                    style.add('BACKGROUND', (5, i), (5, i), colors.red)
                    style.add('TEXTCOLOR', (5, i), (5, i), colors.white)
                elif diff.severity == 'warning':
                    style.add('BACKGROUND', (5, i), (5, i), colors.orange)
            
            detail_table.setStyle(style)
            story.append(detail_table)
        
        doc.build(story)
        output.seek(0)
        return output


def main():
    """Main Streamlit application"""
    
    # Header
    st.markdown('<h1 class="main-header">🔧 OLT MOP Automation Tool</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Compare FIO, EWP, and MOP files between two sites</p>', unsafe_allow_html=True)
    
    # Check Tesseract availability
    if not TESSERACT_AVAILABLE:
        st.warning("⚠️ Tesseract OCR is not installed. EWP image processing will use fallback mode (filename parsing only). For full OCR capabilities, install Tesseract OCR.")
    
    # Sidebar
    with st.sidebar:
        st.image("https://img.icons8.com/color/96/000000/network.png", width=80)
        st.markdown("## About")
        st.markdown("""
        This tool automates the comparison of:
        - **FIO** (Excel files) - OLT connection data
        - **EWP** (Images) - Network diagrams and configurations
        - **MOP** (Word documents) - Method of Procedure documents
        
        It identifies differences between two sites and generates comprehensive reports.
        """)
        
        st.markdown("## Features")
        st.markdown("""
        - ✅ Excel (FIO) parsing
        - ✅ OCR for images (EWP) - optional
        - ✅ Word document parsing (MOP)
        - ✅ Difference analysis
        - ✅ Excel and PDF reports
        - ✅ Cross-reference validation
        """)
        
        if not TESSERACT_AVAILABLE:
            st.info("💡 **OCR Not Available**\n\nEWP images will be processed using filename metadata only. To enable full OCR:\n\n**Ubuntu/Debian:** `sudo apt-get install tesseract-ocr`\n\n**macOS:** `brew install tesseract`")
        
        st.markdown("## Requirements")
        st.markdown("""
        - **FIO**: Excel (.xlsx, .xls)
        - **EWP**: Images (.png, .jpg, .jpeg, .tiff, .bmp)
        - **MOP**: Word (.docx)
        """)
    
    # Main content
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📍 Site 1")
        with st.container():
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            
            # Use unique keys for widgets and separate state variables
            site1_name = st.text_input("Site Name", value="Site A", key="input_site1_name")
            
            site1_fio = st.file_uploader("FIO File (Excel)", type=['xlsx', 'xls'], key="input_site1_fio")
            if site1_fio:
                st.markdown(f'<span class="file-uploaded">✅ {site1_fio.name} ({site1_fio.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            site1_ewp = st.file_uploader("EWP File (Image)", type=['png', 'jpg', 'jpeg', 'tiff', 'bmp'], key="input_site1_ewp")
            if site1_ewp:
                st.markdown(f'<span class="file-uploaded">✅ {site1_ewp.name} ({site1_ewp.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            site1_mop = st.file_uploader("MOP File (Word)", type=['docx'], key="input_site1_mop")
            if site1_mop:
                st.markdown(f'<span class="file-uploaded">✅ {site1_mop.name} ({site1_mop.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown("### 📍 Site 2")
        with st.container():
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            
            site2_name = st.text_input("Site Name", value="Site B", key="input_site2_name")
            
            site2_fio = st.file_uploader("FIO File (Excel)", type=['xlsx', 'xls'], key="input_site2_fio")
            if site2_fio:
                st.markdown(f'<span class="file-uploaded">✅ {site2_fio.name} ({site2_fio.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            site2_ewp = st.file_uploader("EWP File (Image)", type=['png', 'jpg', 'jpeg', 'tiff', 'bmp'], key="input_site2_ewp")
            if site2_ewp:
                st.markdown(f'<span class="file-uploaded">✅ {site2_ewp.name} ({site2_ewp.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            site2_mop = st.file_uploader("MOP File (Word)", type=['docx'], key="input_site2_mop")
            if site2_mop:
                st.markdown(f'<span class="file-uploaded">✅ {site2_mop.name} ({site2_mop.size/1024:.1f}KB)</span>', unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Analyze button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        analyze_clicked = st.button("🔍 Analyze & Compare", use_container_width=True)
    
    # Process when analyze is clicked
    if analyze_clicked:
        # Check if all files are uploaded
        files_uploaded = all([
            site1_fio is not None,
            site1_ewp is not None,
            site1_mop is not None,
            site2_fio is not None,
            site2_ewp is not None,
            site2_mop is not None
        ])
        
        if not files_uploaded:
            st.error("⚠️ Please upload all required files for both sites")
            
            # Show which files are missing
            missing = []
            if site1_fio is None: missing.append("Site 1 FIO")
            if site1_ewp is None: missing.append("Site 1 EWP")
            if site1_mop is None: missing.append("Site 1 MOP")
            if site2_fio is None: missing.append("Site 2 FIO")
            if site2_ewp is None: missing.append("Site 2 EWP")
            if site2_mop is None: missing.append("Site 2 MOP")
            
            if missing:
                st.warning(f"Missing files: {', '.join(missing)}")
        else:
            with st.spinner("Processing files and analyzing differences..."):
                try:
                    # Initialize processors
                    fio_processor = FIOProcessor()
                    ewp_processor = EWPProcessor()
                    mop_processor = MOPProcessor()
                    analyzer = DifferenceAnalyzer()
                    report_generator = ReportGenerator()
                    
                    # Process Site 1
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("Processing Site 1 FIO...")
                    fio1_data = fio_processor.process_file(site1_fio.getvalue(), site1_fio.name)
                    progress_bar.progress(15)
                    
                    status_text.text("Processing Site 1 EWP...")
                    ewp1_data = ewp_processor.process_file(site1_ewp.getvalue(), site1_ewp.name)
                    progress_bar.progress(30)
                    
                    status_text.text("Processing Site 1 MOP...")
                    mop1_data = mop_processor.process_file(site1_mop.getvalue(), site1_mop.name)
                    progress_bar.progress(45)
                    
                    site1_data = SiteData(
                        site_name=site1_name,
                        fio_data=fio1_data,
                        ewp_data=ewp1_data,
                        mop_data=mop1_data,
                        fio_file=site1_fio.name,
                        ewp_file=site1_ewp.name,
                        mop_file=site1_mop.name
                    )
                    
                    # Process Site 2
                    status_text.text("Processing Site 2 FIO...")
                    fio2_data = fio_processor.process_file(site2_fio.getvalue(), site2_fio.name)
                    progress_bar.progress(60)
                    
                    status_text.text("Processing Site 2 EWP...")
                    ewp2_data = ewp_processor.process_file(site2_ewp.getvalue(), site2_ewp.name)
                    progress_bar.progress(75)
                    
                    status_text.text("Processing Site 2 MOP...")
                    mop2_data = mop_processor.process_file(site2_mop.getvalue(), site2_mop.name)
                    progress_bar.progress(90)
                    
                    site2_data = SiteData(
                        site_name=site2_name,
                        fio_data=fio2_data,
                        ewp_data=ewp2_data,
                        mop_data=mop2_data,
                        fio_file=site2_fio.name,
                        ewp_file=site2_ewp.name,
                        mop_file=site2_mop.name
                    )
                    
                    # Analyze differences
                    status_text.text("Analyzing differences...")
                    differences = analyzer.analyze(site1_data, site2_data)
                    summary = report_generator.generate_summary(differences, site1_name, site2_name)
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Analysis complete!")
                    
                    # Store results in session state using unique keys
                    st.session_state['analysis_results'] = {
                        'differences': differences,
                        'summary': summary,
                        'site1_name': site1_name,
                        'site2_name': site2_name,
                        'site1_data': site1_data,
                        'site2_data': site2_data,
                        'analysis_complete': True
                    }
                    
                    st.success("✅ Analysis complete! Scroll down to see results.")
                    
                    if not TESSERACT_AVAILABLE:
                        st.info("ℹ️ Note: EWP images were processed using fallback mode (filename metadata only). Install Tesseract for full OCR capabilities.")
                    
                    # Rerun to display results
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
                    st.exception(e)
    
    # Display results if available
    if 'analysis_results' in st.session_state and st.session_state['analysis_results'].get('analysis_complete', False):
        results = st.session_state['analysis_results']
        differences = results['differences']
        summary = results['summary']
        site1_name = results['site1_name']
        site2_name = results['site2_name']
        
        # Summary metrics
        st.markdown("---")
        st.markdown("## 📊 Analysis Results")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Differences", summary['total_differences'])
        with col2:
            st.metric("Critical", summary['critical_count'], delta_color="inverse")
        with col3:
            st.metric("Warning", summary['warning_count'])
        with col4:
            st.metric("Info", summary['info_count'])
        
        # Category breakdown
        if summary['categories']:
            st.markdown("### Differences by Category")
            cat_df = pd.DataFrame(list(summary['categories'].items()), columns=['Category', 'Count'])
            st.bar_chart(cat_df.set_index('Category'))
        
        # Detailed differences
        st.markdown("### Detailed Differences")
        
        if not differences:
            st.success("🎉 No differences found! Both sites are aligned.")
        else:
            for i, diff in enumerate(differences, 1):
                severity_class = {
                    'critical': 'diff-critical',
                    'warning': 'diff-warning',
                    'info': 'diff-info'
                }.get(diff.severity, 'diff-info')
                
                with st.container():
                    st.markdown(f'<div class="{severity_class}">', unsafe_allow_html=True)
                    col1, col2, col3 = st.columns([1, 3, 1])
                    with col1:
                        st.markdown(f"**#{i}**")
                        st.markdown(f"**{diff.category}**")
                    with col2:
                        st.markdown(f"**{diff.field}**")
                        st.markdown(f"*{diff.description}*")
                    with col3:
                        st.markdown(f"**{diff.severity.upper()}**")
                        st.markdown(f"`{diff.site1_value}` → `{diff.site2_value}`")
                    st.markdown(f"<small>{diff.source1} vs {diff.source2}</small>", unsafe_allow_html=True)
                    st.markdown('</div>', unsafe_allow_html=True)
                    st.markdown("---")
        
        # Data preview
        with st.expander("📁 View Raw Data"):
            tab1, tab2, tab3 = st.tabs(["FIO Data", "EWP Data", "MOP Data"])
            
            with tab1:
                st.markdown(f"**{site1_name} FIO**")
                st.json(results['site1_data'].fio_data)
                st.markdown(f"**{site2_name} FIO**")
                st.json(results['site2_data'].fio_data)
            
            with tab2:
                st.markdown(f"**{site1_name} EWP**")
                st.json(results['site1_data'].ewp_data)
                if not results['site1_data'].ewp_data.get('ocr_available', True):
                    st.info("⚠️ EWP processed in fallback mode (OCR not available)")
                st.markdown(f"**{site2_name} EWP**")
                st.json(results['site2_data'].ewp_data)
            
            with tab3:
                st.markdown(f"**{site1_name} MOP**")
                st.json(results['site1_data'].mop_data)
                st.markdown(f"**{site2_name} MOP**")
                st.json(results['site2_data'].mop_data)
        
        # Report generation
        st.markdown("---")
        st.markdown("## 📄 Generate Reports")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📊 Generate Excel Report", use_container_width=True):
                with st.spinner("Generating Excel report..."):
                    try:
                        report_gen = ReportGenerator()
                        excel_bytes = report_gen.generate_excel_report(
                            differences, site1_name, site2_name
                        )
                        st.download_button(
                            label="📥 Download Excel Report",
                            data=excel_bytes,
                            file_name=f"olt_mop_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                            key="excel_download"
                        )
                        st.success("✅ Excel report ready!")
                    except Exception as e:
                        st.error(f"Error generating Excel report: {e}")
        
        with col2:
            if st.button("📄 Generate PDF Report", use_container_width=True):
                with st.spinner("Generating PDF report..."):
                    try:
                        report_gen = ReportGenerator()
                        pdf_bytes = report_gen.generate_pdf_report(differences, summary)
                        st.download_button(
                            label="📥 Download PDF Report",
                            data=pdf_bytes,
                            file_name=f"olt_mop_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                            key="pdf_download"
                        )
                        st.success("✅ PDF report ready!")
                    except Exception as e:
                        st.error(f"Error generating PDF report: {e}")
    
    # Clear results button
    if 'analysis_results' in st.session_state:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🔄 Clear Results", use_container_width=True):
                st.session_state.pop('analysis_results', None)
                st.rerun()
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 0.8rem;">
        OLT MOP Automation Tool | Built with Streamlit
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
